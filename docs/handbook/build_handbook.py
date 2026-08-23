"""Build Word and PDF versions of the handbook chapters from their markdown.

    python docs/handbook/build_handbook.py            # all chapters, docx + pdf
    python docs/handbook/build_handbook.py 01 02      # only chapters starting 01, 02
    python docs/handbook/build_handbook.py --docx-only

The markdown files are the source of truth. The .docx files exist so that
reviewers can comment in Word; comments are worked back into the markdown and
everything is rebuilt — never hand-edit a generated file. If a .docx is open in
Word it is locked; the build then writes a `-vN` versioned name instead and says
so, leaving the reviewed copy untouched.

Word conversion uses pandoc (on PATH). PDF conversion uses docs/build_pdf.py.
"""

from __future__ import annotations

import argparse
import subprocess
import sys
from pathlib import Path

HERE = Path(__file__).resolve().parent
DOCS = HERE.parent
sys.path.insert(0, str(DOCS))

import build_pdf  # noqa: E402  (docs/build_pdf.py)

OUT_DOCX = HERE / "word"
OUT_PDF = HERE / "pdf"


def versioned(path: Path) -> Path:
    """Next free -vN name if the target is locked or should not be replaced."""
    n = 2
    while True:
        cand = path.with_name("{}-v{}{}".format(path.stem, n, path.suffix))
        if not cand.exists():
            return cand
        n += 1


REVISION_MARKS = {"{{R1}}": (0xB0, 0x00, 0x20)}   # revision round -> RGB (round 1: 2026-08-23)


def color_revisions(docx_path: Path) -> int:
    """Color every paragraph (body or table cell) that carries a revision mark, and strip
    the mark. Pandoc cannot color text from markdown, so this is done on the built file.
    Returns the number of paragraphs marked."""
    import docx
    from docx.shared import RGBColor
    d = docx.Document(str(docx_path))
    paras = list(d.paragraphs)
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                paras.extend(cell.paragraphs)
    n = 0
    for para in paras:
        for mark, rgb in REVISION_MARKS.items():
            if mark in para.text:
                for run in para.runs:
                    run.text = run.text.replace(mark, "")
                    run.font.color.rgb = RGBColor(*rgb)
                n += 1
    d.save(str(docx_path))
    return n


def build_docx(md: Path) -> Path:
    OUT_DOCX.mkdir(exist_ok=True)
    out = OUT_DOCX / (md.stem + ".docx")
    cmd = ["pandoc", str(md), "-o", str(out), "--from", "markdown",
           "--resource-path", str(HERE)]
    try:
        subprocess.run(cmd, check=True, capture_output=True, text=True)
    except subprocess.CalledProcessError as e:
        if out.exists():
            # Probably locked by Word; write a versioned copy instead.
            out = versioned(out)
            subprocess.run(cmd[:3] + [str(out)] + cmd[4:], check=True,
                           capture_output=True, text=True)
            print("  locked -> wrote", out.name)
        else:
            print(e.stderr)
            raise
    n = color_revisions(out)
    if n:
        print("  {} revision-marked paragraphs colored".format(n))
    return out


def build_pdf_file(md: Path) -> Path:
    OUT_PDF.mkdir(exist_ok=True)
    out = OUT_PDF / (md.stem + ".pdf")
    footer = md.stem.replace("_", " ")
    try:
        build_pdf.build(str(md), str(out), footer)
    except PermissionError:
        out = versioned(out)
        build_pdf.build(str(md), str(out), footer)
        print("  locked -> wrote", out.name)
    return out


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("prefixes", nargs="*", help="chapter number prefixes, e.g. 01 02")
    ap.add_argument("--docx-only", action="store_true")
    ap.add_argument("--pdf-only", action="store_true")
    args = ap.parse_args()

    chapters = sorted(HERE.glob("[0-9][0-9]_*.md"))
    if args.prefixes:
        chapters = [c for c in chapters
                    if any(c.name.startswith(p) for p in args.prefixes)]
    if not chapters:
        print("no chapters matched")
        return
    for md in chapters:
        print(md.name)
        if not args.pdf_only:
            print("  docx:", build_docx(md).name)
        if not args.docx_only:
            print("  pdf: ", build_pdf_file(md).name)


if __name__ == "__main__":
    main()
